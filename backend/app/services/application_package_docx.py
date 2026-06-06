"""DOCX export for application packages (markdown → Word, in-memory)."""
from __future__ import annotations

import io
import re
import zipfile

from typing import Literal, Protocol

from docx import Document


DocxPart = Literal["resume", "cover_letter", "strategy"]


class _PackageExport(Protocol):
    resume_markdown: str
    cover_letter_markdown: str
    strategy_notes: str
    job_id: object
    version: int


def markdown_to_docx_bytes(markdown_text: str) -> bytes:
    """Lightweight markdown-like structure: headings, bullets, paragraphs."""
    doc = Document()
    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        if not line:
            doc.add_paragraph("")
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            continue
        if line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
            continue
        doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _sanitize_filename_stub(job_id: object, version: int, tail: str) -> str:
    base = f"job_{job_id}_package_v{version}_{tail}"
    return re.sub(r'[<>:"/\\|?*]', "", base.replace(":", ""))


def single_application_package_docx(pkg: _PackageExport, part: DocxPart) -> tuple[bytes, str]:
    """One ``.docx`` from stored markdown; inner names match the ZIP export."""
    if part == "resume":
        md, tail = pkg.resume_markdown, "resume_draft.docx"
    elif part == "cover_letter":
        md, tail = pkg.cover_letter_markdown, "cover_letter_draft.docx"
    else:
        md, tail = pkg.strategy_notes, "strategy_notes.docx"
    return markdown_to_docx_bytes(md), _sanitize_filename_stub(pkg.job_id, pkg.version, tail)


def zip_application_package_docx(pkg: _PackageExport) -> tuple[bytes, str]:
    """Three .docx files in one zip (`resume`, cover letter, strategy)."""
    bio = io.BytesIO()
    with zipfile.ZipFile(bio, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("resume_draft.docx", markdown_to_docx_bytes(pkg.resume_markdown))
        zf.writestr(
            "cover_letter_draft.docx",
            markdown_to_docx_bytes(pkg.cover_letter_markdown),
        )
        zf.writestr("strategy_notes.docx", markdown_to_docx_bytes(pkg.strategy_notes))
    bio.seek(0)
    name = (
        f"job_{pkg.job_id}_package_v{pkg.version}_docx.zip"
        .replace(":", "")
    )
    return bio.read(), name
