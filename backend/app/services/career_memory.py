"""Personal career memory ingestion and discovery profile (ported from Jobr).

All queries scoped by ``user_id`` (seeded tenant until auth).
Canonical pipeline jobs referenced by UUID via ``canonical_job_id``.
"""
from __future__ import annotations

import io
import json
import logging
import os
import re
import shutil
import subprocess
import sys
import tempfile
import uuid
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from pypdf import PdfReader
from striprtf.striprtf import rtf_to_text
from sqlalchemy import or_, select
from sqlalchemy.orm import Session

from ..config import Settings
from ..models.base import utcnow
from ..models.career_memory import (
    CareerDiscoveryProfile,
    CareerDocument,
    CareerEvidenceChunk,
    CareerFact,
    CareerProfileQuestion,
    CareerTimelineEntry,
)
from ..models.job import Job
from ..services.ai import AIProviderMisconfigured, get_chat_completer

logger = logging.getLogger(__name__)

ROLE_KEYWORDS = {
    "product": ("product manager", "product", "roadmap", "go-to-market", "growth"),
    "operations": ("operations", "program", "process", "cross-functional", "execution"),
    "community": ("community", "devrel", "advocacy", "content", "ecosystem"),
    "marketing": ("marketing", "lifecycle", "campaign", "brand", "demand generation"),
    "data": ("analytics", "data", "sql", "experimentation", "insights"),
}
ADJACENCY_DEFAULTS = {
    "product": ("product marketing", "program manager", "operations manager"),
    "operations": ("chief of staff", "strategy", "program manager"),
    "community": ("developer relations", "partnerships", "ecosystem growth"),
    "marketing": ("growth", "community", "product marketing"),
    "data": ("analytics engineer", "growth analyst", "revops"),
}
SENIORITY_KEYWORDS = ("senior", "lead", "staff", "principal", "director", "head")

_DOCX_CT = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_PDF_CT = "application/pdf"
_DOC_CT = "application/msword"

# Shown when extension/MIME is not one of the supported uploads.
CAREER_MEMORY_UPLOAD_HINT = (
    "Supported files: PDF (.pdf), Word (.docx), legacy Word (.doc), or plain text (.txt). "
    "If .doc conversion fails, install LibreOffice or resave as .docx or PDF."
)


def _strip_nul_chars(s: str) -> str:
    """PostgreSQL text/varchar rejects U+0000; strip defensively before persistence."""
    return s.replace("\x00", "")


def _norm_mime(content_type: str | None) -> str:
    return (content_type or "").split(";")[0].strip().lower()


def _is_pdf_magic(data: bytes) -> bool:
    if len(data) < 4:
        return False
    return data.lstrip()[:4] == b"%PDF"


def _is_zip_magic(data: bytes) -> bool:
    return len(data) >= 4 and data[0:2] == b"PK"


def _is_ole_cfb_magic(data: bytes) -> bool:
    """Compound File Binary (legacy .doc, .xls, …)."""
    return len(data) >= 8 and data[:8] == b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"


def _is_rtf(data: bytes) -> bool:
    head = data[:800].lstrip()
    return head[:5].lower() == b"{\\rtf"


def _find_libreoffice_soffice() -> str | None:
    which = shutil.which("soffice")
    if which:
        return which
    env = (os.environ.get("ATLAS_SOFFICE_PATH") or "").strip()
    if env and Path(env).is_file():
        return env
    if sys.platform == "win32":
        for base in (
            os.environ.get("PROGRAMFILES", r"C:\Program Files"),
            os.environ.get("PROGRAMFILES(X86)", r"C:\Program Files (x86)"),
        ):
            cand = Path(base) / "LibreOffice" / "program" / "soffice.exe"
            if cand.is_file():
                return str(cand)
    return None


def _subprocess_run_quiet(cmd: list[str], *, cwd: str, timeout: int = 120) -> None:
    kwargs: dict = {
        "check": True,
        "timeout": timeout,
        "cwd": cwd,
        "capture_output": True,
        "text": True,
        "encoding": "utf-8",
        "errors": "replace",
    }
    if sys.platform == "win32":
        # Avoid console flash on Windows
        kwargs["creationflags"] = subprocess.CREATE_NO_WINDOW  # type: ignore[attr-defined]
    subprocess.run(cmd, **kwargs)


def _extract_pdf_plain_text(data: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:
        raise ValueError(f"Could not open PDF: {exc}") from exc
    if getattr(reader, "is_encrypted", False):
        raise ValueError("PDF is password-protected; remove the password and re-upload.")
    parts: list[str] = []
    for page in reader.pages:
        t = (page.extract_text() or "").strip()
        if t:
            parts.append(t)
    return "\n\n".join(parts)


def _extract_docx_plain_text(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    parts.append(t)
    return "\n".join(parts)


def _extract_legacy_word_text(data: bytes) -> str:
    """Legacy Word (.doc): RTF body, mislabeled OOXML, or OLE via LibreOffice."""
    if _is_rtf(data):
        decoded = data.decode("cp1252", errors="replace")
        text = rtf_to_text(decoded)
        return text.strip()
    if _is_zip_magic(data):
        try:
            return _extract_docx_plain_text(data).strip()
        except Exception:
            pass
    soffice = _find_libreoffice_soffice()
    if not soffice:
        raise ValueError(
            "Could not read this .doc file. Install LibreOffice, or set ATLAS_SOFFICE_PATH to "
            f"soffice / soffice.exe, or resave as .docx. {CAREER_MEMORY_UPLOAD_HINT}"
        )
    with tempfile.TemporaryDirectory() as td:
        tmp_path = Path(td)
        doc_path = tmp_path / "upload.doc"
        doc_path.write_bytes(data)
        try:
            _subprocess_run_quiet(
                [
                    soffice,
                    "--headless",
                    "--convert-to",
                    "txt:Text",
                    str(doc_path),
                    "--outdir",
                    str(tmp_path),
                ],
                cwd=str(tmp_path),
            )
        except subprocess.CalledProcessError as exc:
            err = (exc.stderr or exc.stdout or str(exc)).strip()
            raise ValueError(f"LibreOffice could not convert this .doc: {err}") from exc
        out = tmp_path / "upload.txt"
        if not out.is_file():
            raise ValueError("LibreOffice did not produce text for this .doc; try resaving as .docx.")
        return out.read_text(encoding="utf-8", errors="replace").strip()


def _smells_like_binary(data: bytes) -> bool:
    sample = data[: min(len(data), 16_384)]
    if not sample:
        return False
    nul = sample.count(b"\x00")
    if nul > max(8, len(sample) // 200):
        return True
    ctrl = sum(1 for b in sample if b < 9 or (b > 13 and b < 32))
    return ctrl > len(sample) * 0.20


def _resolve_upload_kind(filename: str, content_type: str | None, data: bytes) -> str:
    lower = (filename or "upload").lower()
    ct = _norm_mime(content_type)

    if lower.endswith(".pdf") or ct == _PDF_CT:
        return "pdf"
    if lower.endswith(".docx") or ct == _DOCX_CT.lower():
        return "docx"
    if lower.endswith(".doc") or ct == _DOC_CT:
        return "doc"
    if lower.endswith(".txt") or ct.startswith("text/"):
        return "txt"

    if _is_pdf_magic(data):
        return "pdf"
    if _is_zip_magic(data):
        return "docx"
    if _is_ole_cfb_magic(data):
        return "doc"
    if _smells_like_binary(data):
        raise ValueError(CAREER_MEMORY_UPLOAD_HINT)
    return "txt"


def prepare_upload_bytes_as_text(filename: str, content_type: str | None, data: bytes) -> str:
    """Extract plain text from supported uploads: PDF, DOCX, legacy DOC, TXT."""
    if not data:
        raise ValueError("Empty file.")

    kind = _resolve_upload_kind(filename, content_type, data)

    try:
        if kind == "pdf":
            text = _extract_pdf_plain_text(data)
            if not text.strip():
                raise ValueError(
                    "No extractable text in this PDF. If it is a scanned image, use OCR or paste plain text."
                )
            return _strip_nul_chars(text)
        if kind == "docx":
            try:
                text = _extract_docx_plain_text(data)
            except Exception as exc:
                msg = str(exc).strip() or exc.__class__.__name__
                raise ValueError(f"Could not read Word document (.docx): {msg}") from exc
            if not text.strip():
                raise ValueError("Word document appears to have no extractable text.")
            return _strip_nul_chars(text)
        if kind == "doc":
            try:
                text = _extract_legacy_word_text(data)
            except ValueError:
                raise
            except Exception as exc:
                msg = str(exc).strip() or exc.__class__.__name__
                raise ValueError(f"Could not read Word document (.doc): {msg}") from exc
            if not text.strip():
                raise ValueError("No extractable text from this Word document.")
            return _strip_nul_chars(text)
        text = data.decode("utf-8", errors="replace")
        return _strip_nul_chars(text)
    except ValueError:
        raise
    except Exception as exc:
        msg = str(exc).strip() or exc.__class__.__name__
        raise ValueError(msg) from exc


def _chunk_text(text: str, chunk_size: int = 900) -> list[str]:
    return [text[i : i + chunk_size] for i in range(0, len(text), chunk_size)] or [""]


def _extract_candidate_facts(text: str) -> list[str]:
    lines = [line.strip("- ").strip() for line in text.splitlines() if line.strip()]
    filtered = [line for line in lines if len(line.split()) >= 5 and _is_readable_text(line)]
    return filtered[:25]


def _parse_llm_facts_response(raw: str, *, max_items: int) -> list[str]:
    text = (raw or "").strip()
    if not text:
        return []
    if text.startswith("```"):
        lines = text.split("\n")
        if lines and lines[0].strip().startswith("```"):
            lines = lines[1:]
        if lines and lines[-1].strip() == "```":
            lines = lines[:-1]
        text = "\n".join(lines).strip()
    try:
        data = json.loads(text)
    except json.JSONDecodeError:
        start = text.find("{")
        end = text.rfind("}")
        if start >= 0 and end > start:
            try:
                data = json.loads(text[start : end + 1])
            except json.JSONDecodeError:
                return []
        else:
            return []

    items: list = []
    if isinstance(data, dict):
        items = data.get("facts") or data.get("items") or []
    elif isinstance(data, list):
        items = data
    else:
        return []

    out: list[str] = []
    seen: set[str] = set()
    for x in items:
        if not isinstance(x, str):
            continue
        t = _strip_nul_chars(x).strip()
        if len(t) < 12 or len(t) > 4000:
            continue
        key = t.lower()
        if key in seen:
            continue
        seen.add(key)
        out.append(t)
        if len(out) >= max_items:
            break
    return out


def extract_candidate_facts_llm(text: str, *, settings: Settings, document_name: str) -> list[str]:
    """Call OpenAI (via ``get_chat_completer``) to propose atomic draft facts."""
    max_in = int(settings.career_memory_llm_facts_max_input_chars)
    max_items = int(settings.career_memory_llm_facts_max_items)
    body = (text or "").strip()
    truncated = False
    if len(body) > max_in:
        body = body[:max_in]
        truncated = True
    system = (
        "You extract career profile facts from documents for a job-search assistant.\n"
        "Return ONLY valid JSON with this shape: {\"facts\": [\"...\", ...]}.\n"
        "Be EXHAUSTIVE — extract every distinct, verifiable fact. Do NOT summarise multiple facts into one.\n"
        "\n"
        "Extract ALL of the following categories:\n"
        "- ROLES: Each job/role as its own fact including company, title, and date range "
        "(e.g. 'Community Manager at peaq Network, March 2023 to present')\n"
        "- METRICS: Every number, stat, or quantified outcome as its own fact "
        "(e.g. 'Grew Discord to 150k members', 'Newsletter open rate 36-40% for 11,000 subscribers')\n"
        "- RESPONSIBILITIES: Each distinct task, responsibility, or project as its own fact\n"
        "- TOOLS & PLATFORMS: Each named tool, platform, or technology as its own fact or small group "
        "(e.g. 'Experienced with Shopify, Webflow, WordPress, and WooCommerce for ecommerce')\n"
        "- SKILLS: Each distinct skill area as its own fact\n"
        "- ACHIEVEMENTS: Each specific achievement, award, or milestone as its own fact\n"
        "- LIVE EVENTS: Each notable event or broadcast with attendance figures if given\n"
        "- CONTACT/LOCATION: Name, location, email, phone if present\n"
        "\n"
        "Rules:\n"
        "- Split compound facts into separate items — one fact per bullet.\n"
        "- DO include role titles and company names — these are critical facts, not headers.\n"
        "- DO include tool and platform lists as facts.\n"
        "- DO include positioning statements that describe real experience.\n"
        "- Skip pure boilerplate instructions, writing-style notes, and content not about this person.\n"
        "- No duplicates.\n"
        f"- At most {max_items} facts.\n"
        "- Keep each string under 500 characters.\n"
    )
    user_parts = [f"Source label: {document_name}"]
    if truncated:
        user_parts.append("[Note: excerpt was truncated for model context length.]")
    user_parts.append("---")
    user_parts.append(body)
    completer = get_chat_completer(settings)
    reply = completer.complete(
        [
            {"role": "system", "content": system},
            {"role": "user", "content": "\n".join(user_parts)},
        ],
        temperature=0.15,
    )
    return _parse_llm_facts_response(reply, max_items=max_items)


def _parse_year(value: str | None) -> int | None:
    if not value:
        return None
    match = re.search(r"(19|20)\d{2}", value)
    return int(match.group(0)) if match else None


def _guess_date_range(raw: str) -> tuple[str | None, str | None]:
    if not raw:
        return None, None
    cleaned = raw.strip()
    cleaned = re.sub(r"\s+", " ", cleaned)
    parts = re.split(r"\s*(?:-|–|—|to)\s*", cleaned)
    if len(parts) < 2:
        return cleaned[:20], None
    start = parts[0].strip()[:20] if parts[0].strip() else None
    end = parts[1].strip()[:20] if parts[1].strip() else None
    return start, end


def _extract_timeline_entries(cleaned_text: str) -> list[dict]:
    lines = [ln.strip() for ln in cleaned_text.splitlines() if ln.strip()]
    entries: list[dict] = []
    i = 0
    header_patterns = [
        re.compile(
            r"^(?P<title>.+?)\s+(?:at|@)\s+(?P<company>.+?)\s*(?:-|–|—)\s*(?P<dates>.+)$",
            re.I,
        ),
        re.compile(
            r"^(?P<company>.+?)\s*(?:-|–|—)\s*(?P<title>.+?)\s*(?:-|–|—)\s*(?P<dates>.+)$",
            re.I,
        ),
    ]
    while i < len(lines):
        line = lines[i]
        match = None
        for pat in header_patterns:
            match = pat.match(line)
            if match:
                break
        if not match:
            i += 1
            continue

        title = (match.group("title") or "").strip()[:255]
        company = (match.group("company") or "").strip()[:255] or None
        dates_raw = (match.group("dates") or "").strip()
        start_date, end_date = _guess_date_range(dates_raw)
        summary_lines: list[str] = []
        j = i + 1
        while j < len(lines) and len(summary_lines) < 12:
            nxt = lines[j]
            if any(p.match(nxt) for p in header_patterns):
                break
            if len(nxt) >= 8 and _is_readable_text(nxt):
                summary_lines.append(nxt.strip("- ").strip())
            j += 1
        summary = "\n".join(summary_lines).strip() or None
        if title and (_parse_year(dates_raw) or start_date or end_date):
            entries.append(
                {
                    "title": title,
                    "company": company,
                    "start_date": start_date,
                    "end_date": end_date,
                    "summary": summary,
                }
            )
        i = j
    return entries[:25]


def _dates_overlap(a_start: str | None, a_end: str | None, b_start: str | None, b_end: str | None) -> bool:
    y_now = datetime.now(timezone.utc).year
    a1 = _parse_year(a_start)
    a2 = _parse_year(a_end) or y_now
    b1 = _parse_year(b_start)
    b2 = _parse_year(b_end) or y_now
    if a1 is None or b1 is None:
        return False
    return not (a2 < b1 or b2 < a1)


def _maybe_create_timeline_conflict(
    db: Session,
    *,
    user_id: uuid.UUID,
    new_entry: CareerTimelineEntry,
    existing_entry: CareerTimelineEntry,
    source_label: str,
) -> None:
    if (new_entry.company or "").strip().lower() != (existing_entry.company or "").strip().lower():
        return
    if not _dates_overlap(
        new_entry.start_date,
        new_entry.end_date,
        existing_entry.start_date,
        existing_entry.end_date,
    ):
        return
    if (new_entry.title or "").strip().lower() == (existing_entry.title or "").strip().lower():
        return

    stamp = utcnow().isoformat()
    group = (
        new_entry.conflict_group
        or existing_entry.conflict_group
        or f"timeline:{(new_entry.company or 'unknown')[:40]}:{stamp}"
    )
    new_entry.conflict_group = group
    existing_entry.conflict_group = group
    new_entry.status = "conflict"
    existing_entry.status = "conflict"

    question_text = (
        f"Conflict detected in your profile timeline for {(new_entry.company or 'a company')}. "
        f"Which is correct?\n\n"
        f"- Source A ({source_label}): {new_entry.title} ({new_entry.start_date or '?'} - {new_entry.end_date or '?'})\n"
        f"- Source B: {existing_entry.title} ({existing_entry.start_date or '?'} - {existing_entry.end_date or '?'})\n\n"
        "Reply with the correct title and dates (and any missing context)."
    )
    existing_open = db.scalar(
        select(CareerProfileQuestion.id).where(
            CareerProfileQuestion.user_id == user_id,
            CareerProfileQuestion.question_text == question_text,
            CareerProfileQuestion.status == "open",
        ).limit(1)
    )
    if not existing_open:
        db.add(
            CareerProfileQuestion(
                user_id=user_id,
                canonical_job_id=None,
                question_text=question_text,
                question_type="conflict",
                priority="high",
                status="open",
            )
        )


def _is_readable_text(value: str) -> bool:
    if not value or len(value.strip()) < 15:
        return False
    printable_chars = sum(1 for ch in value if ch.isprintable())
    printable_ratio = printable_chars / max(len(value), 1)
    alpha_chars = sum(1 for ch in value if ch.isalpha())
    alpha_ratio = alpha_chars / max(len(value), 1)
    if printable_ratio < 0.9:
        return False
    if alpha_ratio < 0.35:
        return False
    weird_tokens = sum(1 for ch in value if ord(ch) < 32 and ch not in ("\n", "\t", "\r"))
    return weird_tokens == 0


def ingest_source_document(
    db: Session,
    name: str,
    content_type: str | None,
    raw_text: str,
    *,
    user_id: uuid.UUID,
    use_llm_facts: bool = False,
    settings: Settings | None = None,
) -> CareerDocument:
    raw_text = _strip_nul_chars(raw_text)
    cleaned_text = "\n".join(
        line for line in raw_text.splitlines() if _is_readable_text(line) or len(line.split()) >= 3
    )
    doc = CareerDocument(
        user_id=user_id,
        name=name,
        content_type=content_type,
        raw_text=raw_text,
    )
    db.add(doc)
    db.flush()

    chunks = _chunk_text(cleaned_text or raw_text)
    for idx, chunk in enumerate(chunks):
        db.add(CareerEvidenceChunk(source_document_id=doc.id, chunk_text=chunk, chunk_index=idx))

    if use_llm_facts:
        from ..config import get_settings

        cfg = settings or get_settings()
        try:
            get_chat_completer(cfg)
        except AIProviderMisconfigured as exc:
            raise ValueError(
                "LLM fact extraction requires ATLAS_OPENAI_API_KEY in backend environment."
            ) from exc
        fact_texts = extract_candidate_facts_llm(
            cleaned_text or raw_text,
            settings=cfg,
            document_name=name,
        )
        if not fact_texts:
            logger.warning("LLM returned no facts for document %s; using heuristic fallback", name)
            fact_texts = _extract_candidate_facts(cleaned_text or raw_text)
        trace_prefix = "llm"
    else:
        fact_texts = _extract_candidate_facts(cleaned_text or raw_text)
        trace_prefix = "document"

    for fact_text in fact_texts:
        fact_type = "metric" if re.search(r"\d", fact_text) else "experience"
        db.add(
            CareerFact(
                user_id=user_id,
                source_document_id=doc.id,
                fact_text=fact_text,
                fact_type=fact_type,
                verification_state="draft",
                confidence_score=0.72 if fact_type == "metric" else 0.68,
                source_trace=f"{trace_prefix}:{name}",
            )
        )

    extracted_entries = _extract_timeline_entries(cleaned_text or raw_text)
    for row in extracted_entries:
        entry = CareerTimelineEntry(
            user_id=user_id,
            source_document_id=doc.id,
            title=row["title"],
            company=row.get("company"),
            start_date=row.get("start_date"),
            end_date=row.get("end_date"),
            summary=row.get("summary"),
            status="draft",
            confidence_score=0.62,
            source_trace=f"document:{doc.name}",
        )
        db.add(entry)
        db.flush()

        if entry.company:
            existing_candidates = list(
                db.scalars(
                    select(CareerTimelineEntry).where(
                        CareerTimelineEntry.user_id == user_id,
                        CareerTimelineEntry.id != entry.id,
                        CareerTimelineEntry.company.is_not(None),
                    )
                )
            )
            for existing in existing_candidates[:80]:
                _maybe_create_timeline_conflict(
                    db,
                    user_id=user_id,
                    new_entry=entry,
                    existing_entry=existing,
                    source_label=doc.name,
                )

    db.commit()
    db.refresh(doc)
    return doc


def reextract_facts_for_document(
    db: Session,
    document_id: int,
    user_id: uuid.UUID,
    *,
    settings: Settings | None = None,
) -> int:
    """Re-run LLM fact extraction on an existing document.

    Deletes all draft facts linked to the document, runs fresh extraction,
    and inserts new draft facts. Approved and rejected facts are untouched.

    Returns the number of new facts created.
    """
    from ..config import get_settings

    cfg = settings or get_settings()
    doc = db.get(CareerDocument, document_id)
    if not doc or doc.user_id != user_id:
        raise ValueError(f"Document {document_id} not found.")
    if not doc.raw_text:
        raise ValueError("Document has no stored text to re-extract from.")

    # Delete existing draft facts for this doc only
    db.query(CareerFact).filter(
        CareerFact.source_document_id == document_id,
        CareerFact.user_id == user_id,
        CareerFact.verification_state == "draft",
    ).delete(synchronize_session=False)

    cleaned_text = "\n".join(
        line for line in doc.raw_text.splitlines() if _is_readable_text(line) or len(line.split()) >= 3
    )

    try:
        get_chat_completer(cfg)
    except AIProviderMisconfigured as exc:
        raise ValueError(
            "LLM fact extraction requires ATLAS_OPENAI_API_KEY in backend environment."
        ) from exc

    fact_texts = extract_candidate_facts_llm(
        cleaned_text or doc.raw_text,
        settings=cfg,
        document_name=doc.name,
    )
    if not fact_texts:
        logger.warning("LLM returned no facts for document %s on re-extract; using heuristic fallback", doc.name)
        fact_texts = _extract_candidate_facts(cleaned_text or doc.raw_text)

    for fact_text in fact_texts:
        fact_type = "metric" if re.search(r"\d", fact_text) else "experience"
        db.add(
            CareerFact(
                user_id=user_id,
                source_document_id=doc.id,
                fact_text=fact_text,
                fact_type=fact_type,
                verification_state="draft",
                confidence_score=0.72 if fact_type == "metric" else 0.68,
                source_trace=f"llm_reextract:{doc.name}",
            )
        )

    db.commit()
    logger.info("reextract_facts_for_document: %d new facts for doc %d", len(fact_texts), document_id)
    return len(fact_texts)

    db.commit()
    db.refresh(doc)
    return doc


def ingest_reference_document(
    db: Session,
    name: str,
    content_type: str | None,
    raw_text: str,
    *,
    user_id: uuid.UUID,
) -> CareerDocument:
    raw_text = _strip_nul_chars(raw_text)
    doc = CareerDocument(
        user_id=user_id,
        name=name,
        content_type=content_type,
        raw_text=raw_text,
    )
    db.add(doc)
    db.flush()

    cleaned_text = "\n".join(
        line for line in raw_text.splitlines() if _is_readable_text(line) or len(line.split()) >= 3
    )
    chunks = _chunk_text(cleaned_text or raw_text)
    for idx, chunk in enumerate(chunks):
        db.add(CareerEvidenceChunk(source_document_id=doc.id, chunk_text=chunk, chunk_index=idx))

    db.commit()
    db.refresh(doc)
    return doc


def _job_reference_document_name(job: Job) -> str:
    url = (job.apply_url or "").strip()
    suf = f" | {url}" if url else ""
    return (
        f"JOB_POSTING {job.id}: {(job.title or 'Untitled').strip()} "
        f"@ {(job.company_name or 'Unknown').strip()}{suf}"
    )


def ensure_job_posting_reference_document(
    db: Session, job_id: uuid.UUID, *, user_id: uuid.UUID
) -> CareerDocument | None:
    job = db.get(Job, job_id)
    if not job:
        return None
    name = _job_reference_document_name(job)
    existing = db.scalar(
        select(CareerDocument).where(CareerDocument.user_id == user_id, CareerDocument.name == name).limit(1)
    )
    if existing:
        return existing

    raw_text = "\n".join(
        [
            f"Title: {job.title}",
            f"Company: {job.company_name}",
            f"Location: {job.location or ''}".strip(),
            f"Apply URL: {job.apply_url}".strip(),
            "",
            "=== DESCRIPTION ===",
            (job.description_clean or "").strip(),
        ]
    ).strip()
    return ingest_reference_document(db, name=name, content_type="text/plain", raw_text=raw_text, user_id=user_id)


def generate_profile_questions(
    db: Session,
    *,
    user_id: uuid.UUID,
    limit: int = 8,
    canonical_job_id: uuid.UUID | None = None,
) -> list[CareerProfileQuestion]:
    questions: list[CareerProfileQuestion] = []
    existing_open = set(
        db.scalars(
            select(CareerProfileQuestion.question_text).where(
                CareerProfileQuestion.user_id == user_id,
                CareerProfileQuestion.status == "open",
            )
        ).all()
    )

    if canonical_job_id:
        ensure_job_posting_reference_document(db, canonical_job_id, user_id=user_id)

    metric_drafts = list(
        db.scalars(
            select(CareerFact).where(
                CareerFact.user_id == user_id,
                CareerFact.verification_state.in_(["draft", "approved"]),
                CareerFact.fact_type != "metric",
            ).limit(25)
        )
    )
    for fact in metric_drafts[:4]:
        question_text = f"What measurable outcome can you attach to: '{fact.fact_text[:120]}'?"
        if question_text in existing_open:
            continue
        q = CareerProfileQuestion(
            user_id=user_id,
            canonical_job_id=canonical_job_id,
            question_text=question_text,
            question_type="missing_metric",
            priority="high",
            status="open",
        )
        db.add(q)
        questions.append(q)

    weak_facts = list(
        db.scalars(
            select(CareerFact).where(
                CareerFact.user_id == user_id,
                CareerFact.verification_state == "draft",
                CareerFact.confidence_score < 0.6,
            ).limit(20)
        )
    )
    for fact in weak_facts[:4]:
        question_text = f"Can you clarify timeline/scope for: '{fact.fact_text[:120]}'?"
        if question_text in existing_open:
            continue
        q = CareerProfileQuestion(
            user_id=user_id,
            canonical_job_id=canonical_job_id,
            question_text=question_text,
            question_type="clarification",
            priority="medium",
            status="open",
        )
        db.add(q)
        questions.append(q)

    if canonical_job_id:
        job = db.get(Job, canonical_job_id)
        if job:
            job_context = f"{job.title or 'Role'} at {job.company_name or 'Company'}"
            requirement_snippet = (job.description_clean or "")[:180]
            targeted = [
                (
                    f"For {job_context}, which quantified result best proves fit for this requirement: '{requirement_snippet}'?",
                    "job_gap",
                    "high",
                ),
                (
                    f"For {job_context}, what specific project demonstrates the strongest overlap with this posting?",
                    "job_gap",
                    "high",
                ),
                (
                    f"For {job_context}, what objection might a hiring manager raise and how should your evidence answer it?",
                    "job_objection",
                    "medium",
                ),
            ]
            for text_value, q_type, priority in targeted:
                if text_value in existing_open:
                    continue
                q = CareerProfileQuestion(
                    user_id=user_id,
                    canonical_job_id=canonical_job_id,
                    question_text=text_value,
                    question_type=q_type,
                    priority=priority,
                    status="open",
                )
                db.add(q)
                questions.append(q)

    if not questions:
        fallback_questions = [
            ("What are your top 3 measurable outcomes from the last 2 roles?", "missing_metric", "high"),
            ("Which tools/platforms do you want emphasized for target roles?", "positioning", "medium"),
            ("What role types should be deprioritized despite keyword overlap?", "fit_preference", "medium"),
        ]
        for text_value, q_type, priority in fallback_questions:
            if text_value in existing_open:
                continue
            q = CareerProfileQuestion(
                user_id=user_id,
                canonical_job_id=canonical_job_id,
                question_text=text_value,
                question_type=q_type,
                priority=priority,
                status="open",
            )
            db.add(q)
            questions.append(q)

    db.commit()
    return questions[:limit]


def cleanup_unreadable_draft_facts(db: Session, *, user_id: uuid.UUID) -> int:
    draft_facts = list(
        db.scalars(
            select(CareerFact).where(
                CareerFact.user_id == user_id,
                CareerFact.verification_state == "draft",
            )
        )
    )
    rejected = 0
    for fact in draft_facts:
        if not _is_readable_text(fact.fact_text):
            fact.verification_state = "rejected"
            rejected += 1
    db.commit()
    return rejected


def _list_to_csv(values: list[str]) -> str | None:
    cleaned = sorted({v.strip().lower() for v in values if v and v.strip()})
    return ",".join(cleaned) if cleaned else None


def _csv_to_list(value: str | None) -> list[str]:
    if not value:
        return []
    return [item.strip() for item in value.split(",") if item.strip()]


def _derive_profile_components(facts: list[CareerFact]) -> dict:
    fact_text = " ".join((fact.fact_text or "").lower() for fact in facts)
    role_keywords: set[str] = set()
    adjacency_keywords: set[str] = set()
    seniority_keywords: set[str] = set()
    avoid_keywords: set[str] = set()

    for cluster, terms in ROLE_KEYWORDS.items():
        if any(term in fact_text for term in terms):
            role_keywords.update(terms[:3])
            adjacency_keywords.update(ADJACENCY_DEFAULTS.get(cluster, ()))

    for keyword in SENIORITY_KEYWORDS:
        if keyword in fact_text:
            seniority_keywords.add(keyword)

    for fact in facts:
        text = (fact.fact_text or "").lower()
        if "avoid " in text or "not interested" in text or "deprioritize" in text:
            tokens = re.findall(r"[a-zA-Z][a-zA-Z0-9\-\+]{2,}", text)
            avoid_keywords.update(tokens[:5])

    if not role_keywords:
        role_keywords.update({"product", "operations", "growth"})
    if not adjacency_keywords:
        adjacency_keywords.update({"program manager", "strategy", "community"})

    confidence = 0.55 + min(len(facts), 30) * 0.01
    confidence = round(min(confidence, 0.92), 2)
    return {
        "role_keywords": sorted(role_keywords),
        "adjacency_keywords": sorted(adjacency_keywords),
        "seniority_keywords": sorted(seniority_keywords),
        "avoid_keywords": sorted(avoid_keywords),
        "confidence_score": confidence,
    }


def build_or_refresh_discovery_profile(
    db: Session, *, user_id: uuid.UUID, profile_name: str = "default"
) -> CareerDiscoveryProfile:
    facts = list(
        db.scalars(
            select(CareerFact).where(
                CareerFact.user_id == user_id,
                or_(
                    CareerFact.verification_state == "approved",
                    CareerFact.is_core_proof_point == 1,
                ),
            )
        )
    )
    derived = _derive_profile_components(facts)
    profile = db.scalar(
        select(CareerDiscoveryProfile).where(
            CareerDiscoveryProfile.user_id == user_id,
            CareerDiscoveryProfile.profile_name == profile_name,
        )
    )
    if not profile:
        profile = CareerDiscoveryProfile(user_id=user_id, profile_name=profile_name)
        db.add(profile)
    profile.role_keywords_csv = _list_to_csv(derived["role_keywords"])
    profile.adjacency_keywords_csv = _list_to_csv(derived["adjacency_keywords"])
    profile.seniority_keywords_csv = _list_to_csv(derived["seniority_keywords"])
    profile.avoid_keywords_csv = _list_to_csv(derived["avoid_keywords"])
    profile.confidence_score = derived["confidence_score"]
    profile.generated_from_facts = len(facts)
    profile.updated_at = utcnow()
    db.commit()
    db.refresh(profile)
    return profile


def get_discovery_profile_payload(
    db: Session,
    *,
    user_id: uuid.UUID,
    profile_name: str = "default",
    refresh: bool = False,
) -> dict:
    profile = db.scalar(
        select(CareerDiscoveryProfile).where(
            CareerDiscoveryProfile.user_id == user_id,
            CareerDiscoveryProfile.profile_name == profile_name,
        )
    )
    if refresh or not profile:
        profile = build_or_refresh_discovery_profile(db, user_id=user_id, profile_name=profile_name)
    return {
        "profile_name": profile.profile_name,
        "role_keywords": _csv_to_list(profile.role_keywords_csv),
        "adjacency_keywords": _csv_to_list(profile.adjacency_keywords_csv),
        "seniority_keywords": _csv_to_list(profile.seniority_keywords_csv),
        "avoid_keywords": _csv_to_list(profile.avoid_keywords_csv),
        "confidence_score": profile.confidence_score,
        "generated_from_facts": profile.generated_from_facts,
        "updated_at": profile.updated_at.isoformat() if profile.updated_at else None,
    }
